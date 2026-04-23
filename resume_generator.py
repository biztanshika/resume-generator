"""
Resume DOCX Generator
Replicates the VBA macro from vba_code_example.txt using python-docx.
Produces a formatted .docx resume with Century Gothic styling, blue section
headers, bullet lists with bold labels, and a blue header line.
"""


import os
from pathlib import Path


from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml.ns import qn
from lxml import etree




# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
FONT_NAME = "Century Gothic"
BLUE = RGBColor(0x00, 0x66, 0xCC)
LINE_WEIGHT_PT = 2
ASSETS_DIR = Path(__file__).resolve().parent / "assets"


# Maps section title -> icon XML template filename in assets/
SECTION_ICON_MAP = {
    "Experience": "Experience_0.xml",
    "Projects": "Projects_0.xml",
    "Education": "Education_0.xml",
    "Skills": "Skills_0.xml",
}




# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------


def _set_run_font(run, size=11, bold=False, color=None, name=FONT_NAME):
    """Apply font properties to a single Run object."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color




def _make_bullet_paragraph(doc, text, size=11, bold_label=False, indent_level=0):
    """
    Add a bullet paragraph.  If bold_label is True the text before (and
    including) the first colon is set to bold.
    """
    para = doc.add_paragraph(style="List Bullet")
    # Ensure the built-in List Bullet style uses our font
    para.style.font.name = FONT_NAME
    para.style.font.size = Pt(size)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.line_spacing = 1.0


    if indent_level > 0:
        para.paragraph_format.left_indent = Inches(0.25 * indent_level)


    if bold_label and ":" in text:
        colon_pos = text.index(":")
        label = text[: colon_pos + 1]
        rest = text[colon_pos + 1 :]
        run_label = para.add_run(label)
        _set_run_font(run_label, size=size, bold=True)
        run_rest = para.add_run(rest)
        _set_run_font(run_rest, size=size, bold=False)
    else:
        run = para.add_run(text)
        _set_run_font(run, size=size, bold=False)


    return para




def _add_page_break(doc):
    """Insert a page break (new paragraph that starts on the next page)."""
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)
    return para




# ---------------------------------------------------------------------------
# Header blue line  (DrawingML shape injected via XML)
# ---------------------------------------------------------------------------


def add_header_line(doc):
    """
    Add a thin blue horizontal line in the primary header of the first
    section, replicating the VBA Shapes.AddLine call.


    Uses DrawingML <wp:anchor> with page-relative positioning -- the exact
    OOXML equivalent of VBA's:
        .RelativeVerticalPosition   = wdRelativeVerticalPositionPage
        .RelativeHorizontalPosition = wdRelativeHorizontalPositionPage
    """
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False


    page_width_emu = int(section.page_width)        # full page width
    line_y_emu = int(Pt(60))                         # 60 pt from page top = 762000 EMU
    line_weight_emu = int(Pt(2))                     # 2 pt stroke = 25400 EMU


    # Namespace URIs
    W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    A  = "http://schemas.openxmlformats.org/drawingml/2006/main"
    WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"


    # Build a <w:r> containing <w:drawing><wp:anchor>..line shape..
    run_xml = (
        f'<w:r xmlns:w="{W}" xmlns:wp="{WP}" xmlns:a="{A}" xmlns:wps="{WPS}">'
        f'  <w:drawing>'
        f'    <wp:anchor distT="0" distB="0" distL="0" distR="0"'
        f'               simplePos="0" relativeHeight="0" behindDoc="1"'
        f'               locked="1" layoutInCell="1" allowOverlap="1">'
        f'      <wp:simplePos x="0" y="0"/>'
        f'      <wp:positionH relativeFrom="page">'
        f'        <wp:posOffset>0</wp:posOffset>'
        f'      </wp:positionH>'
        f'      <wp:positionV relativeFrom="page">'
        f'        <wp:posOffset>{line_y_emu}</wp:posOffset>'
        f'      </wp:positionV>'
        f'      <wp:extent cx="{page_width_emu}" cy="0"/>'
        f'      <wp:effectExtent l="0" t="0" r="0" b="19050"/>'
        f'      <wp:wrapNone/>'
        f'      <wp:docPr id="1" name="HeaderLine"/>'
        f'      <wp:cNvGraphicFramePr>'
        f'        <a:graphicFrameLocks/>'
        f'      </wp:cNvGraphicFramePr>'
        f'      <a:graphic>'
        f'        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'          <wps:wsp>'
        f'            <wps:cNvCnPr/>'
        f'            <wps:spPr>'
        f'              <a:xfrm>'
        f'                <a:off x="0" y="0"/>'
        f'                <a:ext cx="{page_width_emu}" cy="0"/>'
        f'              </a:xfrm>'
        f'              <a:prstGeom prst="line">'
        f'                <a:avLst/>'
        f'              </a:prstGeom>'
        f'              <a:ln w="{line_weight_emu}">'
        f'                <a:solidFill>'
        f'                  <a:srgbClr val="0066CC"/>'
        f'                </a:solidFill>'
        f'              </a:ln>'
        f'            </wps:spPr>'
        f'            <wps:bodyPr/>'
        f'          </wps:wsp>'
        f'        </a:graphicData>'
        f'      </a:graphic>'
        f'    </wp:anchor>'
        f'  </w:drawing>'
        f'</w:r>'
    )


    run_elem = etree.fromstring(run_xml)


    # Ensure the header has at least one paragraph
    if not header.paragraphs:
        header.add_paragraph()
    hdr_para = header.paragraphs[0]
    hdr_para.text = ""  # clear any default text
    hdr_para._element.append(run_elem)




# ---------------------------------------------------------------------------
# Section helper functions  (1:1 mapping with VBA Subs)
# ---------------------------------------------------------------------------


def insert_name(doc, name):
    """Insert the candidate name (Century Gothic 28pt Bold)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    run = para.add_run(name)
    _set_run_font(run, size=28, bold=True)
    return para




def _ensure_image_rel(doc, image_filename="image1.emf"):
    """
    Add the icon EMF image as a relationship on the document part (if not
    already present). Returns the relationship ID string (e.g. 'rId12').
    """
    image_path = str(ASSETS_DIR / image_filename)


    # Check if this image is already added
    for rel in doc.part.rels.values():
        try:
            if hasattr(rel, 'target_partname') and str(rel.target_partname).endswith(image_filename):
                return rel.rId
            if hasattr(rel, '_target') and hasattr(rel._target, 'partname'):
                if str(rel._target.partname).endswith(image_filename):
                    return rel.rId
        except Exception:
            continue


    # Add a new image part
    from docx.opc.part import Part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    content_type = "image/x-emf"
    with open(image_path, "rb") as f:
        image_blob = f.read()
    image_part = Part(
        partname=PackURI(f"/word/media/{image_filename}"),
        content_type=content_type,
        blob=image_blob,
        package=doc.part.package,
    )
    rid = doc.part.relate_to(image_part, RT.IMAGE)
    return rid




def _add_section_icon(doc, para, section_title):
    """
    Inject the icon anchor XML into the first run of *para* for the given
    section title.  If no icon template exists for the title, does nothing.
    """
    xml_file = SECTION_ICON_MAP.get(section_title)
    if xml_file is None:
        return


    xml_path = ASSETS_DIR / xml_file
    if not xml_path.exists():
        return


    icon_xml = xml_path.read_text(encoding="utf-8")


    # If the icon references the EMF image, ensure we have a relationship
    # and patch the rId in the XML to match the actual one.
    R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    if "r:embed=" in icon_xml or ":embed=" in icon_xml:
        real_rid = _ensure_image_rel(doc)
        # Replace any rIdNN references that point to the image
        import re
        icon_xml = re.sub(r'r:embed="rId\d+"', f'r:embed="{real_rid}"', icon_xml)


    anchor_elem = etree.fromstring(icon_xml.encode("utf-8"))


    # Wrap anchor in <w:drawing> inside a <w:r>
    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    WP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    nsmap = {"w": W}


    run_elem = etree.SubElement(para._element, qn("w:r"), nsmap=nsmap)
    drawing_elem = etree.SubElement(run_elem, qn("w:drawing"))
    drawing_elem.append(anchor_elem)




def insert_section_header(doc, title):
    """
    Century Gothic 16pt Bold Blue section header with an icon, followed
    by a blank paragraph for spacing.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(4)
    run = para.add_run(title)
    _set_run_font(run, size=16, bold=True, color=BLUE)
    # Inject the section icon into this paragraph
    _add_section_icon(doc, para, title)
    # blank spacer
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.space_before = Pt(0)
    _set_run_font(spacer.add_run(""), size=2)
    return para




def insert_summary_inline(doc, items):
    """Bullet list with bold label before the colon (mirrors VBA InsertSummaryInline)."""
    for item in items:
        _make_bullet_paragraph(doc, item, size=11, bold_label=True)




def insert_role_title(doc, title):
    """Job title line: Century Gothic 12pt Bold."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(2)
    run = para.add_run(title)
    _set_run_font(run, size=12, bold=True)
    return para




def insert_technologies_line(doc, tech_text):
    """'Technologies Used: <tech>' with the label portion bold, 11pt."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    label_run = para.add_run("Technologies Used: ")
    _set_run_font(label_run, size=11, bold=True)
    text_run = para.add_run(tech_text)
    _set_run_font(text_run, size=11, bold=False)
    return para




def insert_experience_bullet_list(doc, items):
    """Bulleted duty list (Century Gothic 11pt)."""
    for item in items:
        _make_bullet_paragraph(doc, item, size=11, bold_label=False)




def insert_project_with_number(doc, idx, title, bullets):
    """
    Numbered project heading ('N) Title', bold 11pt) followed by
    bullet list of project details.
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(2)
    run = para.add_run(f"{idx}) {title}")
    _set_run_font(run, size=11, bold=True)
    insert_project_bullet_list(doc, bullets)




def insert_project_bullet_list(doc, items):
    """Bulleted project detail list (Century Gothic 11pt)."""
    for item in items:
        _make_bullet_paragraph(doc, item, size=11, bold_label=False)




def insert_skills_list(doc, items):
    """Skills/certifications bullet list with bold category label."""
    for item in items:
        _make_bullet_paragraph(doc, item, size=11, bold_label=True)




def insert_education_entry(doc, degree, college):
    """Degree (14pt Blue) + college (normal) + blank line."""
    para_deg = doc.add_paragraph()
    para_deg.paragraph_format.space_after = Pt(0)
    para_deg.paragraph_format.space_before = Pt(0)
    run_deg = para_deg.add_run(degree)
    _set_run_font(run_deg, size=14, bold=False, color=BLUE)


    para_col = doc.add_paragraph()
    para_col.paragraph_format.space_after = Pt(0)
    para_col.paragraph_format.space_before = Pt(0)
    run_col = para_col.add_run(college)
    _set_run_font(run_col, size=11, bold=False)


    # trailing blank
    doc.add_paragraph()




# ---------------------------------------------------------------------------
# Section orchestrators
# ---------------------------------------------------------------------------


def build_experience_section(doc, experiences):
    """'Experience' header -> loop through each role."""
    doc.add_paragraph()
    insert_section_header(doc, "Experience")


    for i, exp in enumerate(experiences):
        insert_role_title(doc, exp["role"])
        insert_technologies_line(doc, exp["tech"])
        insert_experience_bullet_list(doc, exp["bullets"])
        # blank separator between roles (except after the last one)
        if i < len(experiences) - 1:
            doc.add_paragraph()




def build_projects_section(doc, projects):
    """'Projects' header -> numbered project entries."""
    doc.add_paragraph()
    insert_section_header(doc, "Projects")


    for idx, proj in enumerate(projects, start=1):
        insert_project_with_number(doc, idx, proj["title"], proj["bullets"])
        # blank separator between projects (except after the last one)
        if idx < len(projects):
            doc.add_paragraph()




# ---------------------------------------------------------------------------
# Main orchestrator
# ---------------------------------------------------------------------------


def create_resume(data, output_path="resume.docx"):
    """
    Build a complete resume DOCX from the supplied data dict and save it.
    """
    doc = Document()


    # -- Document setup: margins, default font --
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.0


    # -- Header blue line --
    add_header_line(doc)


    # -- Name --
    insert_name(doc, data["name"])


    # -- Professional Summary --
    insert_summary_inline(doc, data["summary"])


    # -- Experience --
    build_experience_section(doc, data["experience"])


    # -- Projects --
    build_projects_section(doc, data["projects"])


    # -- Education --
    doc.add_paragraph()
    insert_section_header(doc, "Education")
    education_entries = data["education"]
    if isinstance(education_entries, dict):
        education_entries = [education_entries]
    for edu in education_entries:
        insert_education_entry(doc, edu["degree"], edu["college"])


    # -- Skills --
    insert_section_header(doc, "Skills")
    insert_skills_list(doc, data["skills"])


    # # -- Certifications (optional) --
    # certs = data.get("certifications", [])
    # if certs:
    #     doc.add_paragraph()  # blank line before
    #     insert_section_header(doc, "Certifications")
    #     insert_skills_list(doc, certs)


    # -- Save --
    doc.save(output_path)
    print(f"Resume saved to {output_path}")




# ---------------------------------------------------------------------------
# Sample data  (from vba_code_example.txt  --  John Doe)
# ---------------------------------------------------------------------------


SAMPLE_DATA = {
    "name": "John Doe",
    "summary": [
        "Summary: Results-driven Data Engineer and BI Analyst with 5+ years of experience building data pipelines and analytics solutions.",
        "Data Visualization: Tableau, Power BI",
        "Data Engineering: SQL, Python, Alteryx",
        "BI & Reporting Tools: Tableau Desktop, Power BI, Excel",
        "Infrastructure & Automation: AWS, Docker, ETL workflows",
        "Collaboration & Leadership: Mentoring junior engineers, cross-functional stakeholder engagement",
        "Domain Expertise: Automotive analytics, road network reporting, financial analytics",
        "Problem-Solving: Root cause analysis, performance tuning, process automation",
    ],
    "experience": [
        {
            "role": "Data Engineer",
            "tech": "Alteryx, Tableau, SQL, Python, AWS, Docker",
            "bullets": [
                "Worked as Delivery Engineer for projects with Hyundai, Toyota, Pioneer, BMW, Maruti, Sygic, NRSC.",
                "Built Tableau dashboards for 2.8-4.5% Road network data improvements.",
                "Collaborated with data stewards to standardize datasets to 100% business requirements.",
                "Led code reviews and optimized ETL pipelines, reducing runtime by 30%.",
                "Mentored two junior engineers and interns, driving knowledge transfer and best practices.",
                "Developed automated data pipelines with Python and AWS Lambda to streamline ingestion.",
            ],
        },
        {
            "role": "Business Intelligence Analyst",
            "tech": "Power BI, SQL, Excel, Tableau, Python",
            "bullets": [
                "Designed and deployed interactive Power BI reports for executive stakeholders.",
                "Performed advanced data cleansing and transformation in SQL and Python.",
                "Automated monthly reporting processes, reducing manual effort by 50%.",
                "Collaborated with finance and marketing teams to surface key performance metrics.",
            ],
        },
    ],
    "projects": [
        {
            "title": "E-commerce Website Development",
            "bullets": [
                "Built a responsive and user-friendly e-commerce website using React and Node.js.",
                "Integrated payment gateway and optimized checkout flow.",
            ],
        },
        {
            "title": "Inventory Management System",
            "bullets": [
                "Developed an automated inventory tracking system that reduced manual data entry by 50%.",
                "Implemented notification alerts for low-stock thresholds.",
            ],
        },
    ],
    "education": [
        {
            "degree": "Bachelor of Science in Computer Science",
            "college": "University of California, Los Angeles (UCLA), 2017",
        },
        {
            "degree": "Master of Science in Data Science",
            "college": "Stanford University, 2019",
        },
    ],
    "skills": [
        "Programming: JavaScript, Python, Java",
        "Web Technologies: HTML, CSS, React, Node.js",
        "Databases: MySQL, MongoDB",
        "Tools: Git, Docker, AWS",
        "Methodologies: Agile, Scrum",
    ],
   
}




if __name__ == "__main__":
    create_resume(SAMPLE_DATA, "sample_resume.docx")





