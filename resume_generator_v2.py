"""
Resume Generator v2 -- NEW FORMAT engine.


Builds a .docx using the approved new-format design:


  Thick blue header bar (#3D85C6) repeating on every page
  Name (Century Gothic 28pt, NOT bold, black)
    <thin grey horizontal rule>
  Professional Summary  (16pt bold #366091 + person-silhouette icon)
      prose-style longer-sentence bullets (11pt #333333)
    <thin grey horizontal rule>
  Skills  (16pt bold #366091 + icon)
      "Label: content" bullets (label bold, 11pt #333333)
  Experience  (16pt bold #366091 + icon)
      Role: <name>            (11pt bold)
      Technologies Used: ...  (label bold, 11pt #333333)
      Responsibilities        (11pt bold)
      - bullet points         (11pt #333333)
  Projects  (16pt bold #366091 + icon)
      1) Project Title        (11pt bold)
      - bullet points         (11pt #333333)
  Education  (16pt bold #366091 + icon)
      Degree   (14pt #366091, NOT bold)
      College  (11pt #333333)
  Certifications  (optional, 16pt bold #366091 + icon)


Page: US Letter, 1" margins. 1.25 line spacing.


Public entry point: build_resume(data, output_path).  create_resume is
provided as an alias so callers expecting the v1 signature work unchanged.


This module is consumed by app_v7.py and is fully independent of
resume_generator.py (v1) -- the old engine is not touched.
"""


from pathlib import Path


from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from lxml import etree




# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------
FONT_NAME = "Century Gothic"


NAVY_HEADER = RGBColor(0x36, 0x60, 0x91)   # #366091 -- section headers & degrees
BODY_GREY   = RGBColor(0x33, 0x33, 0x33)   # #333333 -- body text
BLACK       = RGBColor(0x00, 0x00, 0x00)
BAR_BLUE_HEX = "3D85C6"                    # thick top-of-page bar


ASSETS_DIR = Path(__file__).resolve().parent / "assets"


SECTION_ICON_MAP = {
    "Professional Summary": "ProfessionalSummary_0.xml",
    "Experience":     "Experience_0.xml",
    "Projects":       "Projects_0.xml",
    "Education":      "Education_0.xml",
    "Skills":         "Skills_0.xml",
    "Certifications": "Experience_0.xml",   # reuse Experience icon
}


# Image file that each icon-XML asset embeds via r:embed.  XML assets that
# do not reference an image (pure shape groups) can be omitted.
SECTION_ICON_IMAGE = {
    "Professional Summary": "ProfessionalSummary_0.png",
}


# Per-section vertical-offset delta (EMU) to add on top of whatever the icon
# XML already specifies in <wp:positionV><wp:posOffset>.  Positive = shift
# DOWN.  Tuned visually so the icon sits on the same line as the heading text.
SECTION_ICON_V_DELTA_EMU = {
    # Professional Summary XML already carries the approved posV; no shift.
    "Professional Summary": 0,
    "Skills":         90000,
    "Projects":       90000,
    "Education":      90000,
    "Certifications": 90000,
    # "Experience" intentionally omitted -- baseline (delta 0).
}




# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------
def _set_run_font(run, size=11, bold=False, color=None, name=FONT_NAME):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    # also force eastAsia / cs font so Word doesn't fall back
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rFonts.set(qn(f"w:{attr}"), name)




def _tight_paragraph(doc, space_before=0, space_after=0):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.25
    return para




def _add_horizontal_rule(doc, color_hex="BFBFBF", space_before=4, space_after=4):
    """Add a thin horizontal line spanning the page width using a paragraph
    bottom border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    pPr = para._element.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")          # 6 = 0.75 pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    return para




def _make_bullet(doc, text, size=11, bold_label=False, color=BODY_GREY,
                 label_color=None):
    para = doc.add_paragraph(style="List Bullet")
    para.style.font.name = FONT_NAME
    para.style.font.size = Pt(size)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.25


    if bold_label and ":" in text:
        colon = text.index(":")
        label = text[: colon + 1]
        rest = text[colon + 1 :]
        r1 = para.add_run(label)
        _set_run_font(r1, size=size, bold=True,
                      color=label_color if label_color is not None else color)
        r2 = para.add_run(rest)
        _set_run_font(r2, size=size, bold=False, color=color)
    else:
        r = para.add_run(text)
        _set_run_font(r, size=size, bold=False, color=color)
    return para




# ---------------------------------------------------------------------------
# Section icon helpers
# ---------------------------------------------------------------------------
def _ensure_image_rel(doc, image_filename="image1.emf"):
    image_path = ASSETS_DIR / image_filename
    if not image_path.exists():
        return None
    for rel in doc.part.rels.values():
        target = getattr(rel, "target_ref", "") or ""
        if target.endswith(image_filename):
            return rel.rId
    from docx.opc.part import Part
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.packuri import PackURI
    ext = image_path.suffix.lower()
    content_type = {
        ".emf":  "image/x-emf",
        ".png":  "image/png",
        ".jpg":  "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif":  "image/gif",
    }.get(ext, "application/octet-stream")
    blob = image_path.read_bytes()
    part = Part(
        partname=PackURI(f"/word/media/{image_filename}"),
        content_type=content_type,
        blob=blob,
        package=doc.part.package,
    )
    return doc.part.relate_to(part, RT.IMAGE)




def _add_section_icon(doc, para, title):
    """Append the section icon (as an anchored drawing) to the END of the
    heading paragraph.  Icon XML files in assets/ carry their own
    positionH/positionV offsets that place them next to the heading text."""
    xml_file = SECTION_ICON_MAP.get(title)
    if not xml_file:
        return
    xml_path = ASSETS_DIR / xml_file
    if not xml_path.exists():
        return
    icon_xml = xml_path.read_text(encoding="utf-8")


    if "r:embed=" in icon_xml or ":embed=" in icon_xml:
        image_name = SECTION_ICON_IMAGE.get(title, "image1.emf")
        rid = _ensure_image_rel(doc, image_name)
        if rid:
            import re
            icon_xml = re.sub(r'r:embed="rId\d+"', f'r:embed="{rid}"', icon_xml)


    # Optional per-section vertical shift -- patch every
    # <wp:positionV ...><wp:posOffset>N</wp:posOffset> occurrence.
    delta = SECTION_ICON_V_DELTA_EMU.get(title, 0)
    if delta:
        import re
        def _shift(m):
            return f"{m.group(1)}{int(m.group(2)) + delta}{m.group(3)}"
        icon_xml = re.sub(
            r'(<wp:positionV\b[^>]*>\s*<wp:posOffset>)(-?\d+)(</wp:posOffset>)',
            _shift, icon_xml,
        )


    try:
        anchor = etree.fromstring(icon_xml.encode("utf-8"))
    except etree.XMLSyntaxError:
        return


    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    run = etree.SubElement(para._element, qn("w:r"), nsmap={"w": W})
    drawing = etree.SubElement(run, qn("w:drawing"))
    drawing.append(anchor)




def insert_section_header(doc, title, icon=True):
    para = _tight_paragraph(doc, space_before=10, space_after=4)
    r = para.add_run(title)
    _set_run_font(r, size=16, bold=True, color=NAVY_HEADER)
    if icon:
        _add_section_icon(doc, para, title)
    return para




# ---------------------------------------------------------------------------
# Blue header bar (thick filled rectangle near top of every page)
# ---------------------------------------------------------------------------
def add_blue_header_bar(doc, color_hex=BAR_BLUE_HEX,
                        bar_top_emu=274320,   # ~0.30 in from page top
                        bar_height_emu=190500):  # ~0.21 in tall
    """Insert a thick blue horizontal bar in the page header so it repeats
    on every page.


    NOTE on colour: Word's editor view dims the page-header layer when the
    body is being edited, so the bar can look washed out on screen even
    though the stored fill is #3D85C6.  The saved file and every PDF /
    print output show the full saturated colour."""
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False


    page_width_emu = int(section.page_width)
    run_xml = f"""<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                       xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
                       xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
  <w:rPr><w:noProof/></w:rPr>
  <w:drawing>
    <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0"
               relativeHeight="251658240" behindDoc="0" locked="1"
               layoutInCell="1" allowOverlap="1">
      <wp:simplePos x="0" y="0"/>
      <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
      <wp:positionV relativeFrom="page"><wp:posOffset>{bar_top_emu}</wp:posOffset></wp:positionV>
      <wp:extent cx="{page_width_emu}" cy="{bar_height_emu}"/>
      <wp:effectExtent l="0" t="0" r="0" b="0"/>
      <wp:wrapNone/>
      <wp:docPr id="100" name="HeaderBlueBar"/>
      <wp:cNvGraphicFramePr/>
      <a:graphic>
        <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
          <wps:wsp>
            <wps:cNvSpPr/>
            <wps:spPr>
              <a:xfrm>
                <a:off x="0" y="0"/>
                <a:ext cx="{page_width_emu}" cy="{bar_height_emu}"/>
              </a:xfrm>
              <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
              <a:solidFill><a:srgbClr val="{color_hex}"/></a:solidFill>
              <a:ln><a:noFill/></a:ln>
            </wps:spPr>
            <wps:style>
              <a:lnRef idx="0"/>
              <a:fillRef idx="0"/>
              <a:effectRef idx="0"/>
              <a:fontRef idx="minor"/>
            </wps:style>
            <wps:bodyPr/>
          </wps:wsp>
        </a:graphicData>
      </a:graphic>
    </wp:anchor>
  </w:drawing>
</w:r>"""


    run_elem = etree.fromstring(run_xml)
    if not header.paragraphs:
        header.add_paragraph()
    hdr_para = header.paragraphs[0]
    hdr_para.text = ""
    hdr_para._element.append(run_elem)




# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------
def insert_name(doc, name):
    para = _tight_paragraph(doc, space_before=0, space_after=4)
    r = para.add_run(name)
    _set_run_font(r, size=28, bold=False, color=BLACK)
    return para




def insert_summary_bullets(doc, bullets):
    """Prose-style longer-sentence summary bullets (11pt #333333)."""
    for text in bullets:
        _make_bullet(doc, text, size=11, bold_label=False, color=BODY_GREY)




def insert_skills(doc, skills):
    """Bulleted 'Label: content' list (label bold)."""
    for line in skills:
        _make_bullet(doc, line, size=11, bold_label=True,
                     color=BODY_GREY, label_color=BLACK)




def insert_experience(doc, experiences):
    """Each experience -> Role / Technologies Used / Responsibilities + bullets."""
    for idx, exp in enumerate(experiences):
        # Role line
        p = _tight_paragraph(doc, space_before=6 if idx > 0 else 2, space_after=0)
        r_lab = p.add_run("Role: ")
        _set_run_font(r_lab, size=11, bold=True, color=BLACK)
        r_val = p.add_run(exp["role"])
        _set_run_font(r_val, size=11, bold=True, color=BLACK)


        # Technologies Used line
        p = _tight_paragraph(doc, space_before=0, space_after=0)
        r_lab = p.add_run("Technologies Used: ")
        _set_run_font(r_lab, size=11, bold=True, color=BLACK)
        r_val = p.add_run(exp["tech"])
        _set_run_font(r_val, size=11, bold=False, color=BODY_GREY)


        # Responsibilities heading
        p = _tight_paragraph(doc, space_before=2, space_after=0)
        r = p.add_run("Responsibilities")
        _set_run_font(r, size=11, bold=True, color=BLACK)


        # Bulleted responsibilities
        for bullet in exp["bullets"]:
            _make_bullet(doc, bullet, size=11, bold_label=False, color=BODY_GREY)




def insert_projects(doc, projects):
    """Sample-style: '1) Project Title' (11pt bold) followed by bullets."""
    for idx, proj in enumerate(projects, start=1):
        p = _tight_paragraph(doc, space_before=4 if idx > 1 else 2, space_after=0)
        r = p.add_run(f"{idx}) {proj['title']}")
        _set_run_font(r, size=11, bold=True, color=BLACK)
        for bullet in proj["bullets"]:
            _make_bullet(doc, bullet, size=11, bold_label=False, color=BODY_GREY)




def insert_education(doc, entries):
    for edu in entries:
        p = _tight_paragraph(doc, space_before=4, space_after=0)
        r = p.add_run(edu["degree"])
        _set_run_font(r, size=14, bold=False, color=NAVY_HEADER)


        p = _tight_paragraph(doc, space_before=0, space_after=0)
        r = p.add_run(edu["college"])
        _set_run_font(r, size=11, bold=False, color=BODY_GREY)




def insert_certifications(doc, certs):
    for line in certs:
        _make_bullet(doc, line, size=11, bold_label=True, color=BODY_GREY)




# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------
def build_resume(data, output_path):
    """Build the new-format resume from a parsed-JSON dict and save to disk.


    Expected schema (identical to v1):
        {
          "name": str,
          "summary": [str, ...],
          "skills": ["Label: content", ...],
          "experience": [{"role": str, "tech": str, "bullets": [str, ...]}, ...],
          "projects":   [{"title": str, "bullets": [str, ...]}, ...],
          "education":  [{"degree": str, "college": str}, ...],
          "certifications": [str, ...]   # optional
        }
    """
    doc = Document()


    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


    # Normal style defaults
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)


    # Thick blue bar at the top of every page -- lives in the page header
    # so it repeats automatically.
    add_blue_header_bar(doc)


    # ---- Body ----
    insert_name(doc, data["name"])
    _add_horizontal_rule(doc, space_before=2, space_after=6)


    insert_section_header(doc, "Professional Summary", icon=True)
    insert_summary_bullets(doc, data["summary"])
    _add_horizontal_rule(doc, space_before=8, space_after=4)


    insert_section_header(doc, "Skills", icon=True)
    insert_skills(doc, data["skills"])
    _add_horizontal_rule(doc, space_before=8, space_after=4)


    insert_section_header(doc, "Experience", icon=True)
    insert_experience(doc, data["experience"])
    _add_horizontal_rule(doc, space_before=8, space_after=4)


    insert_section_header(doc, "Projects", icon=True)
    insert_projects(doc, data["projects"])
    _add_horizontal_rule(doc, space_before=8, space_after=4)


    insert_section_header(doc, "Education", icon=True)
    edu = data["education"]
    insert_education(doc, edu if isinstance(edu, list) else [edu])


    if data.get("certifications"):
        _add_horizontal_rule(doc, space_before=8, space_after=4)
        insert_section_header(doc, "Certifications", icon=True)
        insert_certifications(doc, data["certifications"])


    doc.save(output_path)




# v1-compatible alias so app_v7 can keep the existing call signature.
create_resume = build_resume



